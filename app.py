import os
import sys
import re
import json
import pytesseract
import fitz
import pythoncom
import spacy
import logging
import pdfplumber
import pandas as pd
import tkinter as tk
from tkinter import filedialog, ttk, messagebox

from win32com import client

from docx import Document
from collections import Counter
from difflib import SequenceMatcher
from concurrent.futures import ThreadPoolExecutor, as_completed

from pdf2image import convert_from_path

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Load Spacy NLP model
if getattr(sys, 'frozen', False):  # Check if running from a packaged executable
    model_path = os.path.join(sys._MEIPASS, 'en_core_web_lg')
else:
    model_path = 'en_core_web_lg'

nlp = spacy.load(model_path)
logging.info("Spacy model loaded successfully.")

# Function to extract text from PDF
def extract_from_pdf(pdf_path):
    logging.info(f"Extracting text from PDF: {pdf_path}")
    first_two_words = ''
    text = ''

    try:

        with pdfplumber.open(pdf_path) as pdf:

            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"


        if len(text)<1:
            doc = fitz.open(pdf_path)
            text = "\n".join([page.get_text("text") for page in doc])

        if len(text) < 1:
            images = convert_from_path(pdf_path)
            text = ''
            for img in images:
                text += pytesseract.image_to_string(img)


        first_two_words = get_first_two_words(text)

    except Exception as e:
        logging.error(f"Error extracting from PDF: {e}")

    #print("text",{text})

    return extract_name_and_email_from_text(text, first_two_words)

# Function to extract text from DOCX
def extract_from_docx(docx_path):
    logging.info(f"Extracting text from DOCX: {docx_path}")
    doc = Document(docx_path)
    text = ''
    first_two_words = ''

    try:
        # Extract text from all paragraphs
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text
            text += para_text + '\n'

            if i == 0:  # Capture the first paragraph's text
                words = para_text.split()
                first_two_words = ' '.join(words[:2])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + '\n'

        for section in doc.sections:
            for header in section.header.paragraphs:
                text += header.text + '\n'
    except Exception as e:
        logging.error(f"Error extracting from DOCX: {e}")

    return extract_name_and_email_from_text(text, first_two_words)


#Extract text from DOC files using pywin32.
def extract_from_doc(doc_path):
    logging.info(f"Extracting text from DOC: {doc_path}")
    doc_path = os.path.normpath(doc_path)
    pythoncom.CoInitialize()  # Initialize COM library for threading support on Windows

    try:
        word = client.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_path)
        text = doc.Content.Text
        doc.Close()
        word.Quit()
    except Exception as e:
        logging.error(f"Error extracting from DOC: {doc_path} {e}")
        return None

    return extract_name_and_email_from_text(text, get_first_two_words(text))


def extract_name_and_email_from_text(text, first_two_words):
    logging.info("Extracting names and emails from text...")

    spacy_names = []
    emails = []
    email_context = ""
    email_line = ""
    first_email = ""
    name_similar_to_email = ""
    name_email_ratio = 0

    try:

        # Apply NLP model to the text
        doc = nlp(text)
        # unique_labels = set(entity.label_ for entity in doc.ents)
        # print("Unique Entity Labels:", unique_labels)

        # Iterate through the detected entities to find names and emails
        for entity in doc.ents:
            if entity.label_ == 'PERSON':
                spacy_names.append(entity.text)

        # use regex to find email

        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{1,}'
        email_match = re.search(email_pattern, text)
        if email_match:
            emails.append(email_match.group(0))

        # Extract first email's surrounding text (3 lines above and below)
        if emails:
            first_email = emails[0]

            lines = text.split("\n")
            email_line_index = next((i for i, line in enumerate(lines) if first_email in line), None)

            if email_line_index is not None:
                start = max(0, email_line_index - 3)  # 3 lines before
                end = min(len(lines), email_line_index + 4)  # 3 lines after
                email_context = "\n".join(lines[start:end])
                email_line = lines[email_line_index]
                context = [first_two_words, email_context, email_line]
                name_similar_to_email, name_email_ratio = find_name_similar_to_email(context, first_email)

    except Exception as e:
        logging.error(f"Error extracting names and emails: {e}")

    return first_two_words, spacy_names, first_email, name_similar_to_email, name_email_ratio



# Function to find name similar to email
def find_name_similar_to_email(sentences, email):
    logging.info("Finding name similar to email...")

    best_match = ""
    best_ratio = 0.0
    email = email.lower()
    target = email.split("@")[0]
    target = re.sub(r'\d+', '', target)
    word_list = []

    try:
        for sentence in sentences:
            word_list.extend(sentence.lower().split())
        # print('words_list',word_list)

        for word in word_list:
            if word == email.lower():
                continue
            ratio = SequenceMatcher(None, target, word).ratio()
            if ratio > best_ratio:  # Keep the best matching word
                best_match = word
                best_ratio = ratio

        for i in range(len(word_list) - 1):
            if '@' in word_list[i] or '@' in word_list[i + 1]:
                continue
            combined_name = f"{word_list[i]}{word_list[i + 1]}"  # Form adjacent two-word name
            combined_name_orig = f"{word_list[i]} {word_list[i + 1]}"
            # print('combined_name1',combined_name)
            ratio = SequenceMatcher(None, target, combined_name).ratio()
            if ratio > best_ratio:
                best_match = combined_name_orig
                best_ratio = ratio
            else:
                combined_name = f"{word_list[i + 1]}{word_list[i]}"  # Form adjacent reverse two-word name
                # print('combined_name2',combined_name)
                ratio = SequenceMatcher(None, target, combined_name).ratio()
                if ratio > best_ratio:
                    best_match = combined_name_orig
                    best_ratio = ratio
    except Exception as e:
        logging.error(f"Error finding name similar to email: {e}")

    return best_match.title(), best_ratio


#Function to get final results
def get_final_results(folder_path):
    logging.info("Generating final results...")
    results = []
    extracted_data = extract_from_all_files(folder_path)
    count = 0
    for filename, first_two_words, spacy_names, email, name_similar_to_email, name_email_ratio in extracted_data:
        logging.info(f"\n\nFile: {filename}")
        logging.info(f"first_two_words: {first_two_words}")
        logging.info(f"Spacy Names: {spacy_names}")
        logging.info(f"Name_similar_to_email: {name_similar_to_email}")
        names = [first_two_words, name_similar_to_email]
        names.extend(spacy_names)
        names = [n.strip().lower().title() for n in names if n.strip()]
        logging.info(f"All names: {names}")
        logging.info(f"unique names {list(set(names))}")

        name_counts = Counter(names)
        selected_name = ""
        if name_counts:
            most_common_name, count = name_counts.most_common(1)[0]
        if count >= 2:
            selected_name = most_common_name
        elif name_similar_to_email.strip() and name_email_ratio > 0.5:
            selected_name = name_similar_to_email
        elif spacy_names:

            best_spacy_ratio=0
            best_spacy=spacy_names[0]
            if len(spacy_names)>1:
                for s_names in spacy_names:
                    ratio = SequenceMatcher(None, email.split("@")[0], s_names).ratio()
                    if ratio> best_spacy_ratio:
                        best_spacy = s_names
                        best_spacy_ratio = ratio
            selected_name = best_spacy
        elif first_two_words.strip():
            selected_name = first_two_words


        logging.info(f"Selected Name: {selected_name}")
        logging.info(f"SelectedEmail: {email}")
        results.append((filename, selected_name, email))

    return results

# Function to process all files in folder
def extract_from_all_files(folder_path):
    logging.info(f"Processing all files in folder: {folder_path}")
    results = []
    file_paths = [
        os.path.join(folder_path, f)
        for f in os.listdir(folder_path)
        if f.lower().endswith((".pdf", ".docx", ".doc"))
    ]

    with ThreadPoolExecutor(max_workers=8) as executor:
        future_to_file = {}
        for path in file_paths:
            if path.lower().endswith(".pdf"):
                future_to_file[executor.submit(extract_from_pdf, path)] = path
            elif path.lower().endswith(".docx"):
                future_to_file[executor.submit(extract_from_docx, path)] = path
            elif path.lower().endswith(".doc"):
                future_to_file[executor.submit(extract_from_doc, path)] = path


        for future in as_completed(future_to_file):
            file_path = future_to_file[future]
            filename = os.path.basename(file_path)
            try:
                first_two_words, names, email, name_similar_to_email, name_email_ratio = future.result()
                results.append((filename, first_two_words, list(set(names)), email, name_similar_to_email, name_email_ratio))
            except Exception as e:
                logging.error(f"Error processing {filename}: {e}")

    return results


# Extract first two words from the entire text
def get_first_two_words(text):
    words = text.strip().split()
    first_two_words = ' '.join(words[:2]) if len(
        words) >= 2 else text.strip()  # Handle cases with <2 words
    return first_two_words


# GUI Application
class CVExtractorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("ITS Name & Email Extractor")
        # Get the screen width and height
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()

        # Set the window width and height
        window_width = 700
        window_height = 400

        # Calculate the position to center the window
        position_top = int(screen_height / 2 - window_height / 2)
        position_left = int(screen_width / 2 - window_width / 2)

        # Set the geometry of the window (x, y, width, height)
        root.geometry(f"{window_width}x{window_height}+{position_left}+{position_top}")

        self.folder_path = tk.StringVar()
        self.results = []
        self.processing=False

        ttk.Label(root, text="Select CV Folder:", font=("Arial", 12)).pack(pady=5)
        frame = ttk.Frame(root)
        frame.pack(pady=5, fill="x", padx=10)

        self.entry_folder = ttk.Entry(frame, textvariable=self.folder_path, width=50)
        self.entry_folder.pack(side="left", padx=5, expand=True, fill="x")
        ttk.Button(frame, text="Browse", command=self.browse_folder).pack(side="left")

        self.processing_label = ttk.Label(root, text="Processing...", font=("Arial", 10), foreground="green")
        self.processing_label.pack(pady=5)
        self.processing_label.pack_forget()

        self.save_excel_button = ttk.Button(root, text="Save to Excel", command=self.save_to_excel)
        self.save_excel_button.pack(pady=5)
        self.save_excel_button.pack_forget()

        self.save_json_button = ttk.Button(root, text="Save to JSON", command=self.save_to_json)
        self.save_json_button.pack(pady=5)
        self.save_json_button.pack_forget()

        tree_frame = ttk.Frame(root)
        tree_frame.pack(pady=10, fill="both", expand=True)

        # Add a vertical scrollbar
        tree_scroll = ttk.Scrollbar(tree_frame, orient="vertical")

        self.tree = ttk.Treeview(
            tree_frame, columns=("Filename", "Name", "Email"),
            show="headings", yscrollcommand=tree_scroll.set
        )

        tree_scroll.config(command=self.tree.yview)
        tree_scroll.pack(side="right", fill="y")
        self.tree.heading("Filename", text="Filename")
        self.tree.heading("Name", text="Name")
        self.tree.heading("Email", text="Email")
        self.tree.pack(pady=10, fill="both", expand=True)

    def browse_folder(self):
        self.processing = True
        self.toggle_processing_label()

        self.results = []
        self.tree.delete(*self.tree.get_children())
        self.toggle_save_button()

        folder = filedialog.askdirectory()


        if folder:
            self.folder_path.set(folder)
            self.process_files()


    def process_files(self):

        folder = self.folder_path.get()
        if not os.path.isdir(folder):
            messagebox.showerror("Error", "Please select a valid folder.")
            return

        self.results = get_final_results(folder)
        #self.tree.delete(*self.tree.get_children())

        for filename, name, email  in self.results:
            self.tree.insert("", "end", values=(filename, name, email))

        self.processing =False
        self.toggle_save_button()
        self.toggle_processing_label()
        messagebox.showinfo("Success", f"Processed {len(self.results)} files.")

    def toggle_save_button(self):
        if len(self.results) > 0:
            self.save_excel_button.pack(pady=5)  # Show the button if there are results
            self.save_json_button.pack(pady=5)
        else:
            self.save_excel_button.pack_forget()
            self.save_json_button.pack_forget()

    def toggle_processing_label(self):
        if self.processing:
            self.processing_label.pack(pady=5)
        else:
            self.processing_label.pack_forget()

    def save_to_excel(self):
        if not self.results:
            messagebox.showerror("Error", "No data to save.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")])
        if not file_path:
            return

        df = pd.DataFrame(self.results, columns=["Filename", "Selected Name", "Email"])
        df.to_excel(file_path, index=False)
        messagebox.showinfo("Success", "Data saved successfully!")


    def save_to_json(self):
        if not self.results:
            messagebox.showerror("Error", "No data to save.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON Files", "*.json")])
        if not file_path:
            return

        # Convert results to a list of dictionaries for better readability in JSON
        data = [{"Filename": filename, "Selected Name": name, "Email": email} for filename, name, email in self.results]

        try:
            with open(file_path, "w", encoding="utf-8") as json_file:
                json.dump(data, json_file, indent=4)
            messagebox.showinfo("Success", "Data saved successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save JSON: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = CVExtractorApp(root)
    root.mainloop()
